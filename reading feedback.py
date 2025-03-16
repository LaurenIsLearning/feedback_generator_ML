import openai
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.exceptions import PackageNotFoundError
from dotenv import load_dotenv

# --- Configuration ---
OLD_FOLDER = "old_with_feedback"  # Contains combined submission+feedback docs
NEW_FOLDER = "new_to_generate"    # Raw student submissions without feedback
OUTPUT_FOLDER = "mirrored_output" # Outputs submissions WITH embedded feedback

# --- Helper Functions ---
def add_feedback_to_doc(original_path, feedback, output_path):
    """Adds feedback to student submission mimicking old example structure"""
    try:
        doc = Document(original_path)
        
        # Add feedback section header (mimic teacher's style)
        doc.add_paragraph("\n\nFeedback:\n").bold = True
        
        # Add feedback content with original formatting
        for paragraph in feedback.split("\n\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.add_run(paragraph.strip())
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error modifying {original_path}: {str(e)}")
        return False

def get_teacher_style():
    """Extracts formatting patterns from old combined docs"""
    for filename in os.listdir(OLD_FOLDER):
        if filename.endswith(".docx"):
            doc = Document(os.path.join(OLD_FOLDER, filename))
            # Analyze structure: usually feedback starts after "Feedback:" heading
            for i, para in enumerate(doc.paragraphs):
                if "Feedback:" in para.text:
                    return {
                        'header_style': para.style,
                        'font': para.runs[0].font.name,
                        'spacing': para.paragraph_format.space_after
                    }
    # Default style if no old docs found
    return {
        'header_style': 'Heading2',
        'font': 'Calibri',
        'spacing': Pt(14)
    }

# --- Modified Main Logic --- 
def main():
    # ... [Previous setup code: folders, OpenAI key, etc.] ...

    # Get formatting style from old examples
    teacher_style = get_teacher_style()

    # Process new submissions
    for filename in os.listdir(NEW_FOLDER):
        if filename.endswith(".docx"):
            # Read student work
            student_path = os.path.join(NEW_FOLDER, filename)
            student_content = read_docx(student_path)
            
            # Generate feedback
            feedback = generate_feedback(teacher_feedback, student_content)
            
            if feedback:
                # Create mirrored document with embedded feedback
                output_path = os.path.join(OUTPUT_FOLDER, filename)
                success = add_feedback_to_doc(student_path, feedback, output_path)
                
                if success:
                    print(f"Created mirrored document: {output_path}")
                else:
                    print(f"Failed to process: {filename}")

if __name__ == "__main__":
    main()
