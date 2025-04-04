from docx import Document
from docx.opc.exceptions import PackageNotFoundError

class BaseExtractor:
    """Shared foundation for both extractors"""
    def __init__(self, file_path):
        self.file_path = file_path
        self.rubric_keywords = ["rubric", "grading", "score", "criteria", "assessment"]

class StudentTextExtractor(BaseExtractor):
    """Core logic from your original StudentTextExtractor"""
    def extract_student_content(self):
        try:
            doc = Document(self.file_path)
            content = "\n".join(
                para.text for para in doc.paragraphs
                if self._is_valid_student_text(para)
            )
            return self._remove_rubric_tables(doc, content)
        except PackageNotFoundError:
            return "Invalid document"
    
    def _is_valid_student_text(self, para):
        text = para.text.strip()
        return (bool(text) and 
               not any(kw in text.lower() for kw in self.rubric_keywords) and
               para.style.name.lower() not in ["title", "heading 1", "heading 2"])

    def _remove_rubric_tables(self, doc, content):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if any(kw in cell.text.lower() for kw in self.rubric_keywords):
                        return content
        return content

class AssignmentTextExtractor(BaseExtractor):
    """Core logic from your original assignment handling"""
    def extract_instructions(self):
        try:
            doc = Document(self.file_path)
            return "\n".join(
                para.text for para in doc.paragraphs
                if self._is_valid_instruction(para)
            )
        except PackageNotFoundError:
            return "Invalid document"
    
    def _is_valid_instruction(self, para):
        text = para.text.strip()
        return (bool(text) and 
               not any(kw in text.lower() for kw in self.rubric_keywords) and
               para.style.name.lower() not in ["title", "header", "footer"])
