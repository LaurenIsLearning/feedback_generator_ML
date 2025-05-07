from docx import Document
from docx.shared import Inches

from tropos.preprocess_docx import StudentSubmission

#---- for future highlight implementation---
#from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Add table borders (python-docx doesn't show them by default)
def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')           # Line thickness
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')     # Black

        tblBorders.append(border)

    tblPr.append(tblBorders)

# Main writer
def write_feedback_to_docx(
    submission_path: str,
    feedback_text: str,
    output_path: str,
    target: StudentSubmission,
    raw_prompt=None,
    debug_output: bool = False,
    include_prompt: bool = False
):
    doc = Document(submission_path)

    # Split the feedback into parts
    inline_text = feedback_text
    summary_text = ""
    rubric_text = ""

    if "--- RUBRIC FEEDBACK ---" in feedback_text:
        inline_summary_part, rubric_text = feedback_text.split("--- RUBRIC FEEDBACK ---", 1)
        if hasattr(target, "rubric") and target.rubric:
            target.rubric.inject_model_feedback(rubric_text)
    else:
        inline_summary_part = feedback_text

    if "--- SUMMARY FEEDBACK ---" in inline_summary_part:
        inline_text, summary_text = inline_summary_part.split("--- SUMMARY FEEDBACK ---", 1)
    else:
        inline_text = inline_summary_part

    # Extract inline feedback pairs with improved parsing
    feedback_pairs = []
    current_quote = None
    current_feedback = None
    
    for line in inline_text.split("\n"):
        line = line.strip()
        if line.startswith("- ") and '"' in line:
            # Save previous pair if exists
            if current_quote and current_feedback:
                feedback_pairs.append((current_quote, current_feedback))
            
            # Start new pair
            parts = line.split('"')
            if len(parts) >= 3:
                current_quote = parts[1]
                current_feedback = parts[2].strip(" -–—:")
            else:
                current_quote = None
                current_feedback = None
        elif current_feedback is not None:
            # Continuation of feedback
            current_feedback += " " + line
    
    # Add the last pair if it exists
    if current_quote and current_feedback:
        feedback_pairs.append((current_quote, current_feedback))

    # Apply inline feedback with improved handling
    for para in doc.paragraphs:
        for quoted, comment in feedback_pairs:
            if quoted in para.text:
                # Clean up the comment by removing any remaining quotes or special markers
                clean_comment = comment.replace('"', '').replace("^^", "'")
                
                parts = para.text.split(quoted)
                if len(parts) >= 2:
                    para.clear()
                    para.add_run(parts[0])
                    bold_run = para.add_run(quoted)
                    bold_run.bold = True
                    italic_run = para.add_run(f" [{clean_comment}]")
                    italic_run.italic = True
                    para.add_run(parts[1])
                break

    # Add summary feedback
    if summary_text.strip():
        doc.add_paragraph("")
        header = doc.add_paragraph("Summary Feedback:")
        header.runs[0].bold = True
        for line in summary_text.strip().split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # Add rubric feedback table (existing code remains the same)
    if target.rubric:
        doc.add_paragraph("")
        rubric_header = doc.add_paragraph("Rubric Feedback:")
        rubric_header.runs[0].bold = True

        table = doc.add_table(rows=1, cols=3)
        set_table_border(table)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Project Portion"
        hdr_cells[1].text = "Ideal Criteria"
        hdr_cells[2].text = "Overall Feedback"

        # Set custom widths
        hdr_cells[0].width = Inches(1.2)
        hdr_cells[1].width = Inches(2.8)
        hdr_cells[2].width = Inches(3.5)

        for portion in target.rubric.get_criteria():
            portion_name = portion["portion"]
            criteria_text = "• " + "\n• ".join(c["text"] for c in portion["criteria"])
            feedback_text = "\n".join(f"- {f['text']}" for f in portion.get("feedback", []))

            row_cells = table.add_row().cells
            row_cells[0].text = portion_name
            row_cells[1].text = criteria_text
            row_cells[2].text = feedback_text

            # Set widths
            row_cells[0].width = Inches(1.2)
            row_cells[1].width = Inches(2.8)
            row_cells[2].width = Inches(3.5)

    # Add debugging information if enabled
    if debug_output:
        doc.add_paragraph("")
        debug_header = doc.add_paragraph("=== Debugging Information ===")
        debug_header.runs[0].bold = True
    
        if include_prompt and include_prompt:
            doc.add_paragraph("")
            prompt_header = doc.add_paragraph("Prompt Used:")
            prompt_header.runs[0].bold = True
            for line in include_prompt.strip().split("\n"):
                doc.add_paragraph(line)
    
        doc.add_paragraph("")
        raw_feedback_header = doc.add_paragraph("Raw Model Feedback:")
        raw_feedback_header.runs[0].bold = True
        for line in feedback_text.strip().split("\n"):
            doc.add_paragraph(line)


    doc.save(output_path)
