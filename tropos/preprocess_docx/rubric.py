from typing import TypedDict, List
from docx import Document
import json
import re


class RubricItem(TypedDict):
    id: str
    text: str


class CriteriaInfo(TypedDict):
    portion: str
    criteria: List[RubricItem]
    feedback: List[RubricItem]


# --------------------------
# Rubric Data Class
# --------------------------
class Rubric:

    criteria: List[RubricItem]
    """ 
    A list that stores all of the criteria for the rubric 
    """

    # NOTE: Not sure why this is here, this should be figured out.
    comments: List[RubricItem]
    """
    All the feedback lists from all of the criteria in one list, for ease of access 
    """

    def __init__(self):
        self.criteria = []  # List of full rubric rows (each as a dictionary)
        self.comments = []  # Just the freeform feedback text (optional)

    def set_criteria(self, criteria):
        self.criteria = criteria
        return self

    def get_criteria(self):
        return self.criteria

    def set_comments(self, comments):
        self.comments = comments
        return self

    def get_comments(self):
        return self.comments

    # -----------------------
    # Formatting for prompt
    # -----------------------
    def format_clean_only(self) -> str:
      """Returns only the ideal rubric criteria, excluding any instructor feedback."""
      text = ""
      for portion in self.criteria:
          text += f"\n== {portion['portion']} ==\n"
          for crit in portion['criteria']:
              text += f"- {crit['text']}\n"
      return text.strip()

    def format_rubric_feedback(self) -> str:
      """Returns past instructor feedback per rubric section."""
      text = ""
      for portion in self.criteria:
            if portion["feedback"]:
                text += f"\n== {portion['portion']} ==\n"
                for fb in portion["feedback"]:
                    text += f"- {fb['text']}\n"
      return text.strip()

    def format_clean_and_feedback(self) -> str:
        """Returns a rubric formatted with ideal criteria AND past instructor feedback per section."""
        text = self.format_clean_only()
        text += self.format_rubric_feedback()
        return text.strip()

    def inject_model_feedback(self, model_text: str):
      """
      Parses GPT-style rubric feedback and injects it into self.criteria.
      Expected format:
      == Section Name ==
      - Feedback comment
      - Feedback comment
      """
      import re
      from collections import defaultdict
  
      rubric_feedback_map = defaultdict(list)
      current_section = None
  
      for line in model_text.strip().split("\n"):
          section_match = re.match(r"==\s*(.*?)\s*==", line)
          comment_match = re.match(r"-\s+(.*)", line)
  
          if section_match:
              current_section = section_match.group(1).strip()
          elif comment_match and current_section:
              rubric_feedback_map[current_section].append(comment_match.group(1).strip())
  
      for section in self.criteria:
          name = section["portion"]
          if name in rubric_feedback_map:
              section["feedback"] = [
                  {"id": f"{name.lower().replace(' ', '_')}_m{i+1}", "text": comment}
                  for i, comment in enumerate(rubric_feedback_map[name])
              ]
  



# --------------------------
# Extracts highlighted text from a cell (ideal criteria column)
# TODO: implement later
# --------------------------


# --------------------------
# Main rubric parser function
# --------------------------
def parse_rubric(doc_path: str) -> "Rubric":
    rubric = Rubric()
    doc = Document(doc_path)

    # Assume last table is rubric
    if not doc.tables:
      # print("No tables in document.")
      return rubric

    table = doc.tables[-1]

    # parse header row
    header_cells = table.rows[0].cells
    headers = [cell.text.strip() for cell in header_cells]

    try:
        # Try to find the column indices by header name
        idx_portion = headers.index("Project Portion")
        idx_criteria = headers.index("Ideal Criteria")
        idx_feedback = headers.index("Overall Feedback")
    except ValueError as e:
        print("Expected rubric column headers not found:", e)
        return rubric

    criteria_data = []
    comments = []

    # Go through each subsequent row in the rubric
    for row in table.rows[1:]:
        cells = row.cells

        portion = cells[idx_portion].text.strip()
        criteria_text = cells[idx_criteria].text.strip()
        feedback = cells[idx_feedback].text.strip()

        # split criteria into individual labeled criteria
        criteria_lines = [
            line.strip() for line in criteria_text.split("\n") if line.strip()
        ]
        criteria_list = [
            {"id": f"{portion.lower().replace(' ', '_')}_c{i+1}", "text": line}
            for i, line in enumerate(criteria_lines)
        ]

        # Separate feedback into individual comments (by newlines or sentence boundaries)
        if "\n" in feedback:
            feedback_parts = [
                part.strip() for part in feedback.split("\n") if part.strip()
            ]
        else:
            feedback_parts = re.split(r"(?<=[.!?])\s+(?=[A-Z])", feedback)
            feedback_parts = [part.strip() for part in feedback_parts if part.strip()]

        feedback_list = [
            {"id": f"{portion.lower().replace(' ', '_')}_f{i+1}", "text": part}
            for i, part in enumerate(feedback_parts)
        ]

        # Store this entire row in a structured format
        criteria_data.append(
            {"portion": portion, "criteria": criteria_list, "feedback": feedback_list}
        )

        comments.extend(feedback_list)  # For quick access later if needed

    # Attach the parsed info to the Rubric object
    rubric.set_criteria(criteria_data)
    rubric.set_comments(comments)

    return rubric

