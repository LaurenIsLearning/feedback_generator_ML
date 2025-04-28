from typing import List
from docx import Document
from docx import table
from docx.table import Table
from docx.text.paragraph import Paragraph


class Submission:
    paragraphs: List[str]

    def __init__(self):
        self._content = ""
        self.paragraphs = []

    def get_content(self) -> str:
        # TODO: Cache this
        return "\n\n".join(self.paragraphs)


def iter_block_items(doc):
    """
    Yield paragraphs and tables in document order.
    """
    parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def parse_submission(doc: Document) -> Submission:
    """Extracts all paragraph text before the final table (rubric), or all text if no table exists."""
    sub = Submission()
    try:
        content = []
        blocks = list(iter_block_items(doc))

        # Find the index of the last table (rubric)
        table_indices = [
            i for i, block in enumerate(blocks) if isinstance(block, Table)
        ]
        if table_indices:
            last_table_index = max(table_indices)
            paragraph_blocks = blocks[:last_table_index]  # cut before rubric
        else:
            paragraph_blocks = blocks  # no rubric, then keep all

        # Collect paragraph text before the last table
        for block in paragraph_blocks:
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    sub.paragraphs.append(text)
                    content.append(text)

    except Exception as e:
        raise RuntimeError(f"Failed to parse submission: {e}")

    return sub
